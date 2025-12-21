import os
import io
import tempfile
import json
import re
import math
import pathlib
from datetime import datetime
from typing import List, Dict, Tuple, Optional

import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from dotenv import load_dotenv
import google.generativeai as genai

# ----------------------
# Configuration & Paths
# ----------------------
# We still use the existing Word template for consistent styling
WORD_TEMPLATE_FILE = "index2/format.docx"

# Model name is hardcoded to match index2-word_converter.py
GEMINI_MODEL_NAME = "models/gemini-2.5-flash"


def load_api_key():
    """Load GOOGLE_API_KEY from .env and configure the Gemini client."""
    load_dotenv()
    api_key = os.getenv("GOOGLE_API_KEY")
    if not api_key:
        raise RuntimeError("GOOGLE_API_KEY not found. Create a .env with GOOGLE_API_KEY=<your_key>.")
    genai.configure(api_key=api_key)


def get_multipage_extraction_prompt():
    """Returns the same detailed prompt as in index2-word_converter.py (unchanged)."""
    return """
    Analyze the entire provided multi-page Marathi PDF document. Find ALL pages that are
    formatted as 'Index-II' (सूची क्र.2).

    **CRITICAL INSTRUCTIONS:**
    1.  **Identify and Process ONLY 'Index-II' pages.** These pages contain numbered
        fields like '(1)विलेखाचा प्रकार', '(5)क्षेत्रफळ', etc.
    2.  **Explicitly IGNORE irrelevant pages.** Completely skip any pages titled
        'Payment Details' or any pages that primarily consist of a table of financial
        transactions. Do not extract any data from them.
    3.  **Return a JSON array.** The final output must be a single JSON array (a list
        of objects), where each JSON object represents the data extracted from ONE
        'Index-II' page.
    4.  If no 'Index-II' pages are found, return an empty array `[]`.

    **For each 'Index-II' page found, extract the following fields into a JSON object:**
    -   `dast_kramank_year`: The year part of 'दस्त क्रमांक'.
    -   `sub_registrar_number`: The number at the end of 'दुय्यम निबंधक'.
    -   `dast_kramank_full`: The full value of 'दस्त क्रमांक'.
    -   `registration_date`: The value for '(10)दस्त नोंदणी केल्याचा दिनांक'.
    -   `document_type`: The value for '(1)विलेखाचा प्रकार'.
    -   `survey_number`: All the numbers inside the double parentheses `((...))` from section (4),
        like 'Survey Number'. Some numbers can have parts to them (for example 1ब, 2ब)
    -   `area_sq_meter`: The value for '(5)क्षेत्रफळ'. VERY IMPORTANT: If the unit is
        'चौ.फुट' (Square Feet) or 'चौ. फूट', convert it to Square Meters by multiplying by
        0.092903. If the unit is 'हेक्टर' (Hectare), convert to Square Meters by multiplying
        by 10,000. If it is already 'चौ.मीटर', use the value directly. Return only the
        final numerical value in Square Meters.
    -   `stamp_duty`: The value for '(12)बाजारभावाप्रमाणे मुद्रांक शुल्क'.
    -   `prakar`: Analyze the text in section '(4) भू-मापन...' and related context. If it contains
        the Marathi word 'सदनिका' or mentions चौ.फूट, set `prakar` to 'सदनिका'. Otherwise set it to
        'बिनशेती जमिन'.
    -   `amount`: The value for '(2) मोबदला'. Return only the numeric value (no currency symbols).

    Return ONLY the JSON array and nothing else.
    """


def clean_and_convert_to_float(value, default=0.0):
    if value is None:
        return default
    try:
        cleaned_value = re.sub(r"[^0-9.]", "", str(value))
        return float(cleaned_value) if cleaned_value else default
    except (ValueError, TypeError):
        return default


# --- Helpers for Survey Number normalization ---

def _is_digit_char(c: str) -> bool:
    return ("0" <= c <= "9") or ("\u0966" <= c <= "\u096F")


def _is_numeric_token(token: str) -> bool:
    token = (token or "").strip()
    if not token:
        return False
    return all(_is_digit_char(ch) for ch in token)


def _starts_with_number(token: str) -> bool:
    token = (token or "").strip()
    return bool(token) and _is_digit_char(token[0])


def normalize_survey_numbers(value) -> str:
    items = None
    if isinstance(value, list):
        items = [str(v).strip() for v in value if str(v).strip()]
    elif isinstance(value, str):
        txt = value.strip()
        if txt.startswith("[") and txt.endswith("]"):
            try:
                parsed = json.loads(txt)
                if isinstance(parsed, list):
                    items = [str(v).strip() for v in parsed if str(v).strip()]
            except Exception:
                items = None
        if items is None:
            if "," in txt:
                items = [s.strip() for s in txt.split(",") if s.strip()]
            else:
                return txt
    else:
        return str(value)

    out = []
    last_base = None
    for tok in items:
        if _is_numeric_token(tok):
            last_base = tok
            out.append(tok)
        elif _starts_with_number(tok) and last_base:
            out.append(f"{last_base} {tok}")
        else:
            out.append(tok)
    return ", ".join(out)


def add_table_borders(table):
    tbl = table._element
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border_el = OxmlElement(f"w:{border_name}")
        border_el.set(qn("w:val"), "single")
        border_el.set(qn("w:sz"), "8")
        border_el.set(qn("w:space"), "0")
        border_el.set(qn("w:color"), "000000")
        tblBorders.append(border_el)
    tbl.tblPr.append(tblBorders)


def _records_from_pdf_bytes(pdf_bytes: bytes) -> List[Dict]:
    """Upload the uploaded PDF (bytes) to Gemini and get a list of records.
    Mirrors process in index2-word_converter.py but works on uploaded file only.
    """
    load_api_key()
    model = genai.GenerativeModel(model_name=GEMINI_MODEL_NAME)

    # Write to a NamedTemporaryFile so genai.upload_file can read it
    with tempfile.NamedTemporaryFile(delete=True, suffix='.pdf') as tmp:
        tmp.write(pdf_bytes)
        tmp.flush()
        pdf_path = pathlib.Path(tmp.name)
        pdf_file = genai.upload_file(path=pdf_path, display_name=pdf_path.name)
        prompt = get_multipage_extraction_prompt()
        response = model.generate_content([prompt, pdf_file])
        json_text = response.text.strip().replace("```json", "").replace("```", "")
        list_of_data = json.loads(json_text)
        if not isinstance(list_of_data, list):
            return []

    all_records = []
    for i, data in enumerate(list_of_data):
        area_sqm = clean_and_convert_to_float(data.get("area_sq_meter"))
        stamp_duty = clean_and_convert_to_float(data.get("stamp_duty"))
        amount = clean_and_convert_to_float(data.get("amount"))
        survey_norm = normalize_survey_numbers(data.get("survey_number"))

        hectares = area_sqm / 10000 if area_sqm > 0 else 0
        rate_per_sqm = stamp_duty / area_sqm if area_sqm > 0 else 0
        rate_per_guntha = rate_per_sqm * 100 if rate_per_sqm > 0 else 0
        rate_per_ha = rate_per_sqm * 10000 if rate_per_sqm > 0 else 0

        processed_data = {
            "dast_kramank_year": data.get("dast_kramank_year", "N/A"),
            "sub_registrar_number": data.get("sub_registrar_number", "N/A"),
            "dast_kramank_full": data.get("dast_kramank_full", "N/A"),
            "registration_date": data.get("registration_date", "N/A"),
            "document_type": data.get("document_type", "N/A"),
            "survey_number": survey_norm if survey_norm else "N/A",
            "area_sq_meter": f"{area_sqm:.4f}",
            "area_hectares": f"{hectares:.8f}",
            "stamp_duty": f"{stamp_duty:.2f}",
            "rate_per_sqm": f"{rate_per_sqm:.2f}",
            "rate_per_guntha": f"{rate_per_guntha:.2f}",
            "rate_per_ha": f"{rate_per_ha:.2f}",
            "prakar": data.get("prakar", "N/A"),
            "amount": f"{amount:.2f}",
            "page_record_num": i + 1,
        }
        all_records.append(processed_data)

    return all_records


def _build_base_table_docx(doc: Document, records: List[Dict]) -> Tuple[List[str], List[List[str]]]:
    """Append rows to template's first table from records. Returns (header, rows) for HTML rendering."""
    # Keep the template's original (Marathi) header labels intact for display
    # Column positions are assumed to match the order we append values below.

    table = doc.tables[0]
    rows_for_html: List[List[str]] = []
    serial_number = 1
    for rec in records:
        values = [
            str(serial_number),                      # 0 व्यवहार क्र. / Serial Number
            rec["dast_kramank_year"],               # 1 Year
            rec["sub_registrar_number"],            # 2 Sub Registrar Number
            rec["dast_kramank_full"],               # 3 Dast Kramank
            rec["registration_date"],               # 4 Registration Date (dd/mm/yyyy expected)
            rec["document_type"],                   # 5 Document Type
            rec["survey_number"],                   # 6 Survey Number
            rec["area_sq_meter"],                   # 7 Area (sq meters)
            rec["area_hectares"],                  # 8 Area (Hectares)
            rec["stamp_duty"],                      # 9 Stamp Duty
            rec["rate_per_sqm"],                    # 10 Rate per SqM
            rec["rate_per_guntha"],                 # 11 Rate per Guntha
            rec["rate_per_ha"],                     # 12 Rate per Ha
            rec.get("prakar", "N/A"),              # 13 प्रकार
            rec.get("amount", "")                  # 14 Amount
        ]
        cells = table.add_row().cells
        for i, val in enumerate(values):
            if i < len(cells):
                cells[i].text = val
        rows_for_html.append(values)
        serial_number += 1

    add_table_borders(table)
    # Read the visual header texts from the template's header row (row index 1)
    try:
        visual_header = [cell.text.strip() for cell in table.rows[1].cells]
    except Exception:
        # Fallback: derive empty headers of same length
        visual_header = [""] * len(table.rows[0].cells)
    return visual_header, rows_for_html


def _to_float(s: str) -> float:
    s_clean = re.sub(r"[^0-9.]", "", s or "")
    try:
        return float(s_clean) if s_clean else 0.0
    except ValueError:
        return 0.0


def _date_in_range(s: str, start_dt: datetime, end_dt: datetime) -> bool:
    try:
        dt = datetime.strptime((s or '').strip(), "%d/%m/%Y")
        return (dt >= start_dt) and (dt <= end_dt)
    except Exception:
        return False


def _build_followup_tables(doc: Document, visual_header: List[str], all_rows: List[List[str]]) -> Dict[str, List[List[str]]]:
    """Re-implements future_filter_and_aggregate() without input(), with temporary rules.
    Returns dict with keys: base_table, filtered_table, derived_table, top_table, and paragraphs list.
    """
    # Column index mapping using the same kept_headers definition
    # Use fixed indices matching the value order inserted in the base table
    idx_serial = 0
    idx_year = 1
    idx_subreg = 2
    idx_dast = 3
    idx_reg_date = 4
    idx_doc_type = 5
    idx_survey = 6
    idx_area_sqm = 7
    idx_area_ha = 8
    idx_stamp = 9
    idx_rate_sqm = 10
    idx_rate_guntha = 11
    idx_rate_ha = 12
    idx_prakar = 13
    idx_amount = 14

    # Data rows start after first 3 rows: title, header, numbers row
    data_rows = all_rows

    # Filter by 'प्रकार' == 'बिनशेती जमिन'
    filtered = [
        r for r in data_rows
        if len(r) > max(idx_prakar, idx_rate_sqm) and r[idx_prakar] == 'बिनशेती जमिन'
    ]

    #! TEMP CHANGE: Use fixed date range instead of input() prompts (will remove later)
    # 14/06/2018 to 13/06/2021
    start_dt = datetime.strptime("14/06/2018", "%d/%m/%Y")
    end_dt = datetime.strptime("13/06/2021", "%d/%m/%Y")
    filtered_in_range = []
    for r in filtered:
        if _date_in_range(r[idx_reg_date], start_dt, end_dt):
            filtered_in_range.append(r)
    filtered = filtered_in_range

    # TEMP CHANGE: remove SN == 8 globally before any downstream computations
    filtered = [r for r in filtered if str(r[idx_serial]).strip() != '8']

    # Sort by '(11) Rate per SqM' desc and keep top 50% (round up)
    filtered.sort(key=lambda r: _to_float(r[idx_rate_sqm]), reverse=True)
    n = len(filtered)
    keep_n = max(1, math.ceil(n * 0.5)) if n > 0 else 0
    top_half = filtered[:keep_n]

    # Insert a heading and a new table for filtered rows (Second table)
    doc.add_paragraph("बिनशेती जमिन - फिल्टर केलेले")
    new_table = doc.add_table(rows=1, cols=len(visual_header))
    new_table.style = doc.tables[0].style
    for c_idx, text in enumerate(visual_header):
        new_table.rows[0].cells[c_idx].text = text

    # Rows already exclude SN 8 above
    for r in filtered:
        cells = new_table.add_row().cells
        for c_idx, text in enumerate(r[:len(visual_header)]):
            cells[c_idx].text = text
    add_table_borders(new_table)

    # Build derived table with selected columns + computed 'दर प्रती चौ.मी.'
    derived_indices = [
        idx_serial,    # Transaction / Serial No.
        idx_survey,    # Survey Number
        idx_area_sqm,  # Area (sq meters)
        idx_area_ha,   # Area (Hectares)
        idx_dast,      # Dast Kramank
        idx_reg_date,  # Registration Date
        idx_amount,    # Amount
        idx_prakar,    # प्रकार
        idx_doc_type   # Document Type
    ]

    doc.add_paragraph("बिनशेती जमिन - निवडक स्तंभ व नवीन 'दर प्रती चौ.मी.' सह")
    derived_header = [visual_header[i] for i in derived_indices] + ["दर प्रती चौ.मी."]
    derived_table = doc.add_table(rows=1, cols=len(derived_header))
    derived_table.style = doc.tables[0].style
    for c_idx, text in enumerate(derived_header):
        derived_table.rows[0].cells[c_idx].text = text

    def compute_rate(row):
        amt = _to_float(row[idx_amount]) if len(row) > idx_amount else 0.0
        area = _to_float(row[idx_area_sqm]) if len(row) > idx_area_sqm else 0.0
        return (amt / area) if area > 0 else 0.0

    derived_rows = []
    for r in filtered:
        rate = compute_rate(r)
        values = [r[i] if i < len(r) else "" for i in derived_indices]
        cells = derived_table.add_row().cells
        for c_idx, text in enumerate(values + [f"{rate:.2f}"]):
            cells[c_idx].text = text
        derived_rows.append((values, rate))
    add_table_borders(derived_table)

    # Third: top 50% of derived by new rate
    derived_rows.sort(key=lambda t: t[1], reverse=True)
    n2 = len(derived_rows)
    keep_n2 = max(1, math.ceil(n2 * 0.5)) if n2 > 0 else 0
    top_half_rows = derived_rows[:keep_n2]

    doc.add_paragraph("बिनशेती जमिन - टॉप 50% (नवीन 'दर प्रती चौ.मी.' नुसार)")
    top_header = derived_header
    top_table = doc.add_table(rows=1, cols=len(top_header))
    top_table.style = doc.tables[0].style
    for c_idx, text in enumerate(top_header):
        top_table.rows[0].cells[c_idx].text = text
    for values, rate in top_half_rows:
        cells = top_table.add_row().cells
        for c_idx, text in enumerate(values + [f"{rate:.2f}"]):
            cells[c_idx].text = text
    add_table_borders(top_table)

    # Average
    avg_value = (sum(rate for _, rate in top_half_rows) / len(top_half_rows)) if top_half_rows else 0.0
    avg_paragraph = f"Average दर प्रती चौ.मी. = {avg_value:.2f}"
    doc.add_paragraph(avg_paragraph)

    # Prepare structures for HTML rendering
    def table_to_rows(t):
        return [[cell.text for cell in row.cells] for row in t.rows]

    return {
        "filtered_table": table_to_rows(new_table),
        "derived_table": table_to_rows(derived_table),
        "top_table": table_to_rows(top_table),
        "avg_paragraph": avg_paragraph,
    }


def _render_tables_as_html(base_header: List[str], base_rows: List[List[str]], followup: Dict[str, List[List[str]]]) -> str:
    """Render Word-like tables into HTML preserving existing page styling.
    Uses class 'data' so current CSS in templates/index.html applies.
    """
    def render_table(rows: List[List[str]]):
        if not rows:
            return ""
        head_html = "<thead><tr>" + "".join(f"<th>{pd.isna(h) and '' or h}</th>" for h in rows[0]) + "</tr></thead>"
        body_rows = rows[1:] if len(rows) > 1 else []
        body_html = "<tbody>" + "".join(
            "<tr>" + "".join(f"<td>{pd.isna(c) and '' or c}</td>" for c in r) + "</tr>" for r in body_rows
        ) + "</tbody>"
        return f"<table class=\"data\">{head_html}{body_html}</table>"

    # Build the base table (header + rows)
    base_rows_all = [base_header] + base_rows
    html_parts = []
    html_parts.append("<h3>Index-II Extracted Records</h3>")
    html_parts.append(render_table(base_rows_all))

    # Follow-up tables
    html_parts.append("<h3>बिनशेती जमिन - फिल्टर केलेले</h3>")
    html_parts.append(render_table(followup.get("filtered_table", [])))

    html_parts.append("<h3>बिनशेती जमिन - निवडक स्तंभ व नवीन 'दर प्रती चौ.मी.' सह</h3>")
    html_parts.append(render_table(followup.get("derived_table", [])))

    html_parts.append("<h3>बिनशेती जमिन - टॉप 50% (नवीन 'दर प्रती चौ.मी.' नुसार)</h3>")
    html_parts.append(render_table(followup.get("top_table", [])))

    html_parts.append(f"<p><strong>{followup.get('avg_paragraph','')}</strong></p>")
    return "\n".join(html_parts)


def process_index2_pdf_to_html(pdf_bytes: bytes) -> Tuple[str, Optional[str]]:
    """Public entry point used by Flask route.
    - Extract records with Gemini
    - Build python-docx document using template (for layout parity)
    - Reproduce the same tables into HTML for the frontend
    - No CSV is written; Word is not saved to disk
    """
    # Extract records
    records = _records_from_pdf_bytes(pdf_bytes)

    # Prepare docx in memory (for structure & styling parity)
    doc = Document(WORD_TEMPLATE_FILE)

    # Base table
    base_header, base_rows = _build_base_table_docx(doc, records)

    # Follow-up tables (no input(), fixed dates, and removing व्यवहार क्र. == 8 in second table)
    followup = _build_followup_tables(doc, visual_header=base_header, all_rows=base_rows)

    # Render to HTML with current page styles
    html = _render_tables_as_html(base_header, base_rows, followup)

    # Save DOCX to temp for download (requested by user)
    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_doc:
            doc.save(tmp_doc.name)
            tmp_path = tmp_doc.name
    except Exception:
        tmp_path = None

    return html, tmp_path
