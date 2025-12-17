# 3) Optional transliteration
try:
    from unidecode import unidecode  # type: ignore
except Exception:
    def unidecode(x: str) -> str:  # fallback no-op
        return x
import re
import time
from io import BytesIO
import sys
from typing import List, Tuple, Optional, Dict, Callable

from docx import Document
from bs4 import BeautifulSoup
from playwright.sync_api import sync_playwright
from urllib.parse import quote_plus

"""Optional translation support"""
# 1) Prefer deep_translator (synchronous)
try:
    from deep_translator import GoogleTranslator as _DeepGoogleTranslator  # type: ignore
    _deep_translator_available = True
except Exception:
    _DeepGoogleTranslator = None
    _deep_translator_available = False

# 2) Fallback: googletrans (typically synchronous in 4.0.0rc1)
try:
    from googletrans import Translator as _GTTranslator  # type: ignore
    _gt_translator = _GTTranslator()
except Exception:
    _gt_translator = None

# ----------------------
# Helpers: parse admin fields and survey numbers from .docx
# ----------------------

LABEL_DISTRICT = ("जिल्हा", "जिल्हा:")
LABEL_TALUKA = ("तालुका", "तालुका:")
LABEL_VILLAGE = ("मौजे", "मौजे:")
ALL_LABEL_TOKENS = ("मौजे", "मौजे:", "तालुका", "तालुका:", "जिल्हा", "जिल्हा:")
TABLE_SURVEY_HEADER = "भूमापन क्रमांक / गट क्रमांक"

# Simple in-memory cache for translations
_translate_cache: Dict[str, str] = {}

def _translate_to_en(name: Optional[str]) -> Optional[str]:
    """Translate arbitrary text to English using googletrans if available.
    If translation is unavailable or fails, return the original text.
    """
    if not name:
        return name
    t = name.strip()
    if not t:
        return t
    if t in _translate_cache:
        return _translate_cache[t]
    # Try deep_translator first (explicit Marathi source)
    if _deep_translator_available:
        try:
            out = _DeepGoogleTranslator(source='mr', target='en').translate(t)
            out = (out or '').strip() or t
            _translate_cache[t] = out
            print(f"[Method2] Translated district via deep_translator '{t}' -> '{out}'")
            return out
        except Exception:
            pass

    # Fallback to googletrans if available
    if _gt_translator is not None:
        try:
            res = _gt_translator.translate(t, src='mr', dest='en')
            out = (getattr(res, 'text', None) or '').strip() or t
            _translate_cache[t] = out
            print(f"[Method2] Translated district via googletrans '{t}' -> '{out}'")
            return out
        except Exception:
            pass

    # Last resort: transliterate to ASCII as a heuristic English form
    try:
        ascii_name = unidecode(t).strip()
        if ascii_name:
            print(f"[Method2] Transliteration fallback '{t}' -> '{ascii_name}'")
            _translate_cache[t] = ascii_name
            return ascii_name
    except Exception:
        pass
    return t

def _clean_text(s: str) -> str:
    # Normalize whitespace (including non-breaking) and trim
    if not s:
        return ""
    s = s.replace("\u200c", "").replace("\u00a0", " ")
    s = re.sub(r"\s+", " ", s)
    return s.strip()


def _truncate_at_next_label(text: str) -> str:
    """Given the remainder after a label, cut off at the next label occurrence if present.
    This handles lines where multiple labels appear in the same paragraph.
    """
    if not text:
        return text
    # Find earliest index of any label token in the remainder (excluding at position 0)
    cut_idx = None
    for tok in ALL_LABEL_TOKENS:
        idx = text.find(tok)
        if idx != -1 and idx != 0:
            cut_idx = idx if cut_idx is None else min(cut_idx, idx)
    if cut_idx is not None:
        return _clean_text(text[:cut_idx])
    return _clean_text(text)


def _extract_field_from_paragraphs(doc: Document, labels: Tuple[str, ...]) -> Optional[str]:
    for p in doc.paragraphs:
        text = _clean_text(p.text)
        for lbl in labels:
            if lbl in text:
                # Take part after label and colon if present
                parts = re.split(r"[:：]", text, maxsplit=1)
                if len(parts) == 2 and lbl in parts[0]:
                    return _truncate_at_next_label(parts[1])
                # Otherwise, take text after label token
                idx = text.find(lbl)
                remainder = text[idx + len(lbl):]
                remainder = remainder.lstrip(':：').strip()
                return _truncate_at_next_label(remainder)
    return None


def _find_table_column_index(doc: Document, header_text: str) -> Tuple[Optional[int], Optional[Document]]:
    header_text_norm = _clean_text(header_text)
    for tbl in doc.tables:
        # Assume first row is header
        if len(tbl.rows) == 0:
            continue
        hdr_cells = tbl.rows[0].cells
        for j, cell in enumerate(hdr_cells):
            if header_text_norm in _clean_text(cell.text):
                return j, tbl
    return None, None


def _consider_survey_number(raw: str) -> Optional[str]:
    # Examples:
    # 123 -> 123
    # 123/3 -> 123
    # 123/A -> 123A
    # 123/इ -> 123इ
    # 123/4/अ -> 123
    # 123/क/गट -> 123क
    txt = _clean_text(raw)
    if not txt:
        return None

    def _emit_status(self, msg: str):
        # Simple status hook for UI/console
        try:
            print(f"[Status] {msg}")
        except Exception:
            pass
    # Keep only first 2 segments
    parts = [p.strip() for p in txt.split('/') if p.strip()]
    if not parts:
        return None
    # First must start with digits
    m = re.match(r"^(\d+)", parts[0])
    if not m:
        return None
    base_num = m.group(1)
    if len(parts) == 1:
        return base_num
    second = parts[1]
    # If second starts with a digit => ignore, just the base number
    if re.match(r"^\d+", second):
        return base_num
    # If second starts with a letter (Latin or Devanagari) then append letters from it
    m2 = re.match(r"^([A-Za-z\u0900-\u097F]+)", second)
    if m2:
        return base_num + m2.group(1)
    return base_num


def extract_admin_and_surveys_from_docx(file_bytes: bytes) -> Tuple[Optional[str], Optional[str], Optional[str], List[str]]:
    doc = Document(BytesIO(file_bytes))

    district = _extract_field_from_paragraphs(doc, LABEL_DISTRICT)
    taluka = _extract_field_from_paragraphs(doc, LABEL_TALUKA)
    village = _extract_field_from_paragraphs(doc, LABEL_VILLAGE)

    col_idx, table = _find_table_column_index(doc, TABLE_SURVEY_HEADER)
    surveys: List[str] = []
    if col_idx is not None and table is not None:
        for i in range(1, len(table.rows)):
            cell_text = _clean_text(table.rows[i].cells[col_idx].text)
            if cell_text:
                considered = _consider_survey_number(cell_text)
                if considered:
                    surveys.append(considered)
    return district, taluka, village, surveys


# ----------------------
# Scraper for the IGR site using the new logic
# ----------------------

class IGRSubzoneScraper:
    def __init__(self, headless: bool = False, progress_cb: Optional[Callable[[str], None]] = None):
        self.headless = headless
        self.progress_cb = progress_cb
        self.playwright = None
        self.browser = None
        self.page = None
        self.base_url = 'https://igreval.maharashtra.gov.in/eASR2.0/eASRCommon.aspx?hDistName='
        self._last_survey_input = None  # (ctx, locator) of the survey input we filled

    def __enter__(self):
        print('[Method2] Starting Playwright...')
        self.playwright = sync_playwright().start()
        self.browser = self.playwright.chromium.launch(
            headless=self.headless,
            slow_mo=200,
            args=[
                '--no-sandbox',
                '--disable-dev-shm-usage',
                '--disable-gpu',
                '--disable-extensions',
                '--lang=en-IN'                     # 👈 force Chromium language
            ]
        )

        context = self.browser.new_context(
            locale='en-IN',                        # 👈 navigator.language
            extra_http_headers={
                'Accept-Language': 'en-IN,en;q=0.9'  # 👈 server-side language
            }
        )

        self.page = context.new_page()

        self.page.set_default_timeout(30000)
        return self

    def __exit__(self, exc_type, exc, tb):
        try:
            if self.page:
                print('[Method2] Closing page')
                self.page.close()
        finally:
            try:
                if self.browser:
                    print('[Method2] Closing browser')
                    self.browser.close()
            finally:
                if self.playwright:
                    print('[Method2] Stopping Playwright')
                    self.playwright.stop()

    def _emit_status(self, msg: str):
        """Emit lightweight status updates for UI/console."""
        try:
            print(f"[Status] {msg}")
            try:
                sys.stdout.flush()
            except Exception:
                pass
        except Exception:
            pass
        # Forward to external progress callback if provided (for frontend live updates)
        try:
            if callable(self.progress_cb):
                self.progress_cb(msg)
        except Exception:
            pass

    # --- Utilities ---
    def _find_context_with_selector(self, selector: str, timeout: int = 30000):
        """Return page or first frame that contains the selector."""
        try:
            self.page.wait_for_selector(selector, timeout=timeout)
            return self.page
        except Exception:
            for fr in self.page.frames:
                try:
                    fr.wait_for_selector(selector, timeout=2000)
                    return fr
                except Exception:
                    continue
        return None

    def _select_dropdown_label(self, selector: str, label: str):
        ctx = self._find_context_with_selector(selector)
        if ctx is None:
            raise RuntimeError(f"Selector not found in any frame: {selector}")
        print(f"[Method2] Selecting option '{label}' on {selector}")
        ctx.select_option(selector, label=label)
        time.sleep(1.0)

    def _click_by_text_any(self, text: str, exact: bool = True):
        targets = [self.page] + [f for f in self.page.frames]
        for ctx in targets:
            try:
                loc = ctx.get_by_text(text, exact=exact)
                if loc.count() > 0:
                    print(f"[Method2] Clicking text {'exact' if exact else 'contains'}='{text}'")
                    loc.first.click()
                    return True
            except Exception:
                continue
        return False

    def _log_checked_radio(self, prefix: str = ""):
        targets = [self.page] + [f for f in self.page.frames]
        for ctx in targets:
            try:
                checked = ctx.locator("input[type='radio']:checked")
                if checked.count() > 0:
                    try:
                        rid = checked.first.get_attribute('id')
                        label = ''
                        if rid:
                            lab = ctx.locator(f"label[for='{rid}']")
                            if lab.count() > 0:
                                label = lab.first.inner_text().strip()
                        if not label:
                            label = checked.first.evaluate("el => (el.closest('label')?.innerText) || (el.parentElement?.innerText) || '' ") or ''
                        print(f"[Method2] {prefix}Checked radio label: '{label}'")
                        return
                    except Exception:
                        pass
            except Exception:
                continue

    def _select_radio_option(self, option_label: str) -> bool:
        """Select a radio option by its visible label text.
        Tries clicking the text, then clicking the radio associated with a matching label.
        """
        # Try clicking by visible text
        if self._click_by_text_any(option_label, exact=True):
            return True
        # Try approximate
        if self._click_by_text_any(option_label, exact=False):
            return True
        # Try label->radio association
        if self._click_radio_by_label(option_label):
            return True
        # As a last resort, scan all radios and match keywords in associated label/nearby text
        base = option_label
        keywords = [
            base,
            base.replace('.', ''),
            base.replace('No.', 'No'),
            base.replace(' ', ''),
        ]
        # Add common variants
        if 'Survey' in base:
            keywords += ['SurveyNo', 'SurveyNo.', 'survey', 'surveyno']
        if 'SubZone' in base or 'SubZones' in base:
            keywords += ['Sub Zones', 'Sub Zone', 'subzones', 'sub zone']
        if self._scan_radios_and_click(keywords):
            return True
        return False

    def _ensure_radio_selected(self, option_label: str) -> bool:
        """Force select the radio associated with a visible label (without relying on click).
        Finds a label that contains option_label, then checks the nearest radio input using JS.
        """
        targets = [self.page] + [f for f in self.page.frames]
        for ctx in targets:
            try:
                lab = ctx.locator(f"//label[contains(normalize-space(.), '{option_label}')]")
                if lab.count() == 0:
                    continue
                # Prefer following radio then preceding
                for xpath in [
                    f"//label[contains(normalize-space(.), '{option_label}')]/following::input[@type='radio'][1]",
                    f"//label[contains(normalize-space(.), '{option_label}')]/preceding::input[@type='radio'][1]",
                ]:
                    radio = ctx.locator(xpath)
                    if radio.count() > 0:
                        try:
                            radio.first.evaluate("el => { el.checked = true; el.dispatchEvent(new Event('change', {bubbles:true})); el.dispatchEvent(new Event('click', {bubbles:true})); }")
                            # Verify
                            try:
                                if radio.first.is_checked():
                                    return True
                            except Exception:
                                return True
                        except Exception:
                            continue
            except Exception:
                continue
        return False

    def _click_radio_by_label(self, label_text: str) -> bool:
        """Try to click a radio button associated with the given label text."""
        targets = [self.page] + [f for f in self.page.frames]
        for ctx in targets:
            try:
                # Try label then click the nearest preceding/following radio
                lab = ctx.locator(f"//label[contains(normalize-space(.), '{label_text}')]")
                if lab.count() > 0:
                    # Following input
                    try:
                        radio = ctx.locator(f"//label[contains(normalize-space(.), '{label_text}')]/following::input[@type='radio'][1]")
                        if radio.count() > 0:
                            radio.first.click()
                            return True
                    except Exception:
                        pass
                    # Preceding input
                    try:
                        radio = ctx.locator(f"//label[contains(normalize-space(.), '{label_text}')]/preceding::input[@type='radio'][1]")
                        if radio.count() > 0:
                            radio.first.click()
                            return True
                    except Exception:
                        pass
            except Exception:
                continue
        return False

    def _scan_radios_and_click(self, keywords: List[str]) -> bool:
        """Scan all radio inputs on page/frames and click the one whose associated label or nearby text contains any keyword."""
        keys = [k.strip().lower() for k in keywords if k and k.strip()]
        targets = [self.page] + [f for f in self.page.frames]
        for ctx in targets:
            try:
                radios = ctx.locator("input[type='radio']")
                count = radios.count()
                for i in range(count):
                    r = radios.nth(i)
                    label_text = ''
                    try:
                        rid = r.get_attribute('id')
                        if rid:
                            lab = ctx.locator(f"label[for='{rid}']")
                            if lab.count() > 0:
                                label_text = lab.first.inner_text().strip()
                    except Exception:
                        pass
                    if not label_text:
                        # Try closest following/preceding label
                        try:
                            foll = r.locator("xpath=following::label[1]")
                            if foll.count() > 0:
                                label_text = foll.first.inner_text().strip()
                        except Exception:
                            pass
                    if not label_text:
                        # Try nearby container text
                        try:
                            nearby = r.evaluate("el => (el.closest('td')?.innerText) || (el.closest('tr')?.innerText) || (el.parentElement?.innerText) || '' ")
                            label_text = (nearby or '').strip()
                        except Exception:
                            pass
                    lt = (label_text or '').lower()
                    if any(k in lt for k in keys):
                        try:
                            r.first.click()
                            return True
                        except Exception:
                            continue
                print(f"[Method2] Radio scan did not find keywords: {keys}")
            except Exception:
                continue
        return False

    def _set_input_by_label_text(self, label_text: str, value: str):
        # Strategies to locate the input associated with a textual label
        targets = [self.page] + [f for f in self.page.frames]
        strategies = []
        # 0) Direct known id from probe
        strategies.append(lambda ctx: ctx.locator("#ctl00_ContentPlaceHolder5_txtCommonSurvey"))
        # 1) <label>Enter Survey No</label> -> following input
        strategies.append(lambda ctx: ctx.locator(f"//label[contains(normalize-space(.), '{label_text}')]/following::input[1]"))
        # 2) Any element containing the text -> following input
        strategies.append(lambda ctx: ctx.locator(f"//*[contains(normalize-space(.), '{label_text}')]/following::input[1]"))
        # 3) Input with aria-label/placeholder containing words
        for key in ['survey', 'enter survey', 'survey no']:
            strategies.append(lambda ctx, k=key: ctx.locator(f"//input[contains(translate(@aria-label,'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'), '{k}')]"))
            strategies.append(lambda ctx, k=key: ctx.locator(f"//input[contains(translate(@placeholder,'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'), '{k}')]"))
        # 4) Common id/name contains
        for frag in ['Survey', 'survey', 'txtSurvey', 'txtSearch', 'Search', 'srch']:
            strategies.append(lambda ctx, f=frag: ctx.locator(f"//input[contains(@id, '{f}') and (@type='text' or not(@type))]"))
            strategies.append(lambda ctx, f=frag: ctx.locator(f"//input[contains(@name, '{f}') and (@type='text' or not(@type))]"))
        # 5) Input immediately preceding the Search button
        strategies.append(lambda ctx: ctx.locator("(//input[@type='text'])[last()]"))
        strategies.append(lambda ctx: ctx.locator("(//button[contains(., 'Search')] | //input[@type='submit' and contains(@value,'Search')] | //input[@type='button' and contains(@value,'Search')] | //a[contains(.,'Search')])[1]/preceding::input[@type='text'][1]"))

        for ctx in targets:
            for build in strategies:
                try:
                    loc = build(ctx)
                    if loc and loc.count() > 0:
                        # Ensure element is visible and enabled
                        try:
                            loc.first.wait_for(state='visible', timeout=3000)
                        except Exception:
                            pass
                        print(f"[Method2] Filling survey input found via strategy with value '{value}'")
                        loc.first.fill(value)
                        # Remember exact input used so we can find the correct results table relative to it
                        try:
                            self._last_survey_input = (ctx, loc.first)
                        except Exception:
                            self._last_survey_input = None
                        return True
                except Exception:
                    continue
        raise RuntimeError(f"Input for label '{label_text}' not found")

    def _wait_and_get_results_table_html(self, timeout_ms: int = 12000):
        """Wait for the first results table to appear after Search and return (context, inner_html).
        Heuristics: prefer table with id '#ctl00_ContentPlaceHolder5_ruralDataGrid'. Otherwise, pick
        a table with > 1 body rows and containing expected header keywords.
        """
        import time as _t
        deadline = _t.time() + (timeout_ms / 1000.0)
        targets = [self.page] + [f for f in self.page.frames]
        preferred_sel = '#ctl00_ContentPlaceHolder5_ruralDataGrid'
        header_keywords = ['Attribute', 'ऑद्योगिक', 'उपविभाग', 'Rs.', 'Rate', 'रू.', 'खुली जमीन', 'निवासी सदनिका', 'दुकाने']
        # Function to score a candidate table as Survey results
        def _score_table(html: str) -> int:
            try:
                soup = BeautifulSoup(f"<table>{html}</table>", 'html.parser')
                rows = soup.find_all('tr')
                if len(rows) <= 1:
                    return 0
                text_all = soup.get_text(' ', strip=True)
                # Header score
                hscore = sum(1 for k in header_keywords if k.lower() in text_all.lower())
                # First column pattern rows (start with digit or digit+/)
                valid_rows = 0
                for tr in rows[1:]:
                    tds = tr.find_all('td')
                    if not tds:
                        continue
                    fc = tds[0].get_text(strip=True)
                    if fc and __import__('re').match(r"^\d+[\-/\)]?", fc):
                        valid_rows += 1
                # Prefer small result sets typical of Survey No search (<20)
                small_bonus = 2 if 1 <= valid_rows <= 20 else 0
                return hscore * 3 + valid_rows * 2 + small_bonus
            except Exception:
                return 0
        last_error = None
        # First, try to locate table immediately following the exact survey input we filled
        if self._last_survey_input is not None:
            try:
                ctx, input_loc = self._last_survey_input
                # Check first few following tables to avoid layout container
                best = None
                best_score = 0
                for k in range(1, 6):
                    foll_tbl = input_loc.locator(f"xpath=following::table[{k}]")
                    if not foll_tbl or foll_tbl.count() == 0:
                        continue
                    html = foll_tbl.first.inner_html()
                    if not html or not html.strip():
                        continue
                    # Skip known SubZones pager/grid signature
                    if 'dg_Valuation2_0' in html:
                        continue
                    try:
                        soup = BeautifulSoup(f"<table>{html}</table>", 'html.parser')
                        rows = soup.find_all('tr')
                        if len(rows) <= 1:
                            continue
                        text_all = soup.get_text(" ", strip=True)
                        exclude = ['Select Taluka', 'Select Village', 'Search By', 'Enter Survey No', 'Year', 'Language']
                        if any(x.lower() in text_all.lower() for x in exclude):
                            continue
                        score = _score_table(html)
                        if score > best_score:
                            best = (ctx, html)
                            best_score = score
                    except Exception:
                        continue
                if best is not None and best_score > 0:
                    return best
            except Exception:
                pass
        # If that fails, approximate via geometry from a nearby input
        anchor_rect = None
        try:
            for ctx in targets:
                try:
                    anchor = ctx.locator("//input[@type='text']").first
                    if anchor and anchor.count() > 0:
                        anchor_rect = anchor.evaluate("el => el.getBoundingClientRect()")
                        if anchor_rect:
                            break
                except Exception:
                    continue
        except Exception:
            anchor_rect = None
        while _t.time() < deadline:
            for ctx in targets:
                try:
                    # Preferred selector
                    if ctx.locator(preferred_sel).count() > 0:
                        try:
                            ctx.wait_for_selector(preferred_sel, timeout=1000)
                        except Exception:
                            pass
                        html = ctx.locator(preferred_sel).inner_html()
                        if html and html.strip():
                            return ctx, html
                except Exception as e:
                    last_error = e
                try:
                    # Fallback: any table with >1 data rows
                    tables = ctx.locator('table')
                    tcount = min(tables.count(), 10)  # scan first 10
                    nearest = None
                    nearest_dist = None
                    best_scored = None
                    best_sc = 0
                    for i in range(tcount):
                        tloc = tables.nth(i)
                        thtml = tloc.inner_html()
                        if not thtml or len(thtml) < 30:
                            continue
                        # Skip known SubZones pager/grid signature
                        if 'dg_Valuation2_0' in thtml:
                            continue
                        try:
                            soup = BeautifulSoup(f"<table>{thtml}</table>", 'html.parser')
                            rows = soup.find_all('tr')
                            if len(rows) <= 1:
                                continue
                            head_txt = soup.get_text(" ", strip=True)
                            if any(k.lower() in head_txt.lower() for k in header_keywords):
                                # If we have an anchor, choose the nearest table below the anchor
                                if anchor_rect:
                                    try:
                                        rect = tloc.evaluate("el => el.getBoundingClientRect()")
                                        if rect and rect.get('top', 0) >= anchor_rect.get('top', 0):
                                            dist = rect.get('top', 0) - anchor_rect.get('bottom', 0)
                                            if nearest is None or (dist >= 0 and dist < (nearest_dist or 1e9)):
                                                nearest = (ctx, thtml)
                                                nearest_dist = dist
                                            # also score it
                                            sc = _score_table(thtml)
                                            if sc > best_sc:
                                                best_scored = (ctx, thtml)
                                                best_sc = sc
                                            continue
                                    except Exception:
                                        pass
                                # Fallback: if no anchor or rect issues, return first matching table
                                sc = _score_table(thtml)
                                if sc > best_sc:
                                    best_scored = (ctx, thtml)
                                    best_sc = sc
                        except Exception:
                            continue
                    # Prefer best scored candidate if available; else nearest
                    if best_scored is not None:
                        return best_scored
                    if nearest is not None:
                        return nearest
                except Exception as e:
                    last_error = e
            _t.sleep(0.4)
        if last_error:
            print(f"[Method2] Table wait last error: {last_error}")
        return None, ''

    def _try_select_district_any(self, district_text: str) -> bool:
        """Fallback: scan all <select> elements and try selecting option whose visible text contains the district.
        Tries translated English first, then original.
        """
        candidates = [district_text]
        tr = _translate_to_en(district_text)
        if tr and tr.lower() != district_text.lower():
            candidates.insert(0, tr)
        print(f"[Method2] Fallback district selection candidates: {candidates}")
        targets = [self.page] + [f for f in self.page.frames]
        for ctx in targets:
            try:
                selects = ctx.locator('select')
                count = selects.count()
                for i in range(count):
                    sel = selects.nth(i)
                    # Read options text
                    opt_count = sel.locator('option').count()
                    option_texts = []
                    for j in range(opt_count):
                        option_texts.append(sel.locator('option').nth(j).inner_text().strip())
                    for cand in candidates:
                        for txt in option_texts:
                            if cand.lower() in txt.lower():
                                print(f"[Method2] Selecting district '{txt}' in generic select[{i}]")
                                sel.select_option(label=txt)
                                time.sleep(1.0)
                                return True
            except Exception:
                continue
        return False

    # --- Main flow ---
    def run(self, district: str, year_label: str, taluka: str, village: str, surveys: List[str], translate_admin: bool = True) -> Dict:
        # Navigate with district param (use English name in URL via translation)
        self._emit_status('Navigating IGR')
        district_param = _translate_to_en(district) if translate_admin else district
        district_param = district_param or district
        # Ensure URL-safe
        district_param_enc = quote_plus(district_param)
        url = f"{self.base_url}{district_param_enc}"
        print(f"[Method2] District from doc: '{district}' | using in URL as: '{district_param}' -> encoded '{district_param_enc}'")
        print(f"[Method2] Navigating to: {url}")
        self.page.goto(url, wait_until='domcontentloaded')
        try:
            self.page.wait_for_load_state('networkidle', timeout=30000)
        except Exception:
            pass
        # Page is loaded
        self._emit_status('Inputting Taluka, Village and Year values')

        # Select year; if not found, fallback to base URL and manual district selection
        print('[Method2] Waiting for Year dropdown...')
        try:
            self._select_dropdown_label('#ctl00_ContentPlaceHolder5_ddlYear', year_label)
            # Year selected
            self._emit_status('Inputting Taluka, Village and Year values')
        except Exception as e:
            print(f"[Method2] Year dropdown not found after direct URL. Fallback: reload base and select district. Reason: {e}")
            base_url = 'https://igreval.maharashtra.gov.in/eASR2.0/eASRCommon.aspx'
            self.page.goto(base_url, wait_until='domcontentloaded')
            try:
                self.page.wait_for_load_state('networkidle', timeout=30000)
            except Exception:
                pass
            if not self._try_select_district_any(district):
                return {"error": "Could not select district from page"}
            # After district change, the year dropdown should be present
            self._select_dropdown_label('#ctl00_ContentPlaceHolder5_ddlYear', year_label)
            self._emit_status('Inputting Taluka, Village and Year values')

        # Select taluka
        if translate_admin:
            taluka_en = _translate_to_en(taluka) or taluka
            print(f"[Method2] Selecting taluka: original='{taluka}' | translated='{taluka_en}'")
            try:
                self._select_dropdown_label('#ctl00_ContentPlaceHolder5_ddlTaluka', taluka_en)
            except Exception:
                print("[Method2] Taluka English selection failed, trying original label")
                self._select_dropdown_label('#ctl00_ContentPlaceHolder5_ddlTaluka', taluka)
        else:
            print(f"[Method2] Selecting taluka (no translation): '{taluka}'")
            self._select_dropdown_label('#ctl00_ContentPlaceHolder5_ddlTaluka', taluka)
        time.sleep(1.0)
        # Taluka selected
        self._emit_status('Inputting Taluka, Village and Year values')

        # Wait for village dropdown to be ready
        print('[Method2] Waiting for village dropdown to populate...')
        try:
            self.page.wait_for_function(
                "document.querySelector('#ctl00_ContentPlaceHolder5_ddlVillage') && document.querySelector('#ctl00_ContentPlaceHolder5_ddlVillage').options.length > 1",
                timeout=30000
            )
        except Exception:
            pass
        # Select village
        ctx_village = self._find_context_with_selector('#ctl00_ContentPlaceHolder5_ddlVillage') or self.page
        opts_len = ctx_village.locator('#ctl00_ContentPlaceHolder5_ddlVillage option').count()
        chosen_value = None
        if translate_admin:
            village_key = _translate_to_en(village) or village
            print(f"[Method2] Matching village by translated English: '{village_key}' (original='{village}')")
        else:
            village_key = village
            print(f"[Method2] Matching village (no translation): '{village_key}'")
        for i in range(opts_len):
            opt_text = ctx_village.locator('#ctl00_ContentPlaceHolder5_ddlVillage option').nth(i).inner_text().strip()
            if village_key.lower() in opt_text.lower():
                chosen_value = ctx_village.locator('#ctl00_ContentPlaceHolder5_ddlVillage option').nth(i).get_attribute('value')
                print(f"[Method2] Matched village option '{opt_text}'")
                break
        if chosen_value:
            ctx_village.select_option('#ctl00_ContentPlaceHolder5_ddlVillage', value=chosen_value)
        else:
            # fallback select by label
            try:
                self._select_dropdown_label('#ctl00_ContentPlaceHolder5_ddlVillage', village_key)
            except Exception:
                self._select_dropdown_label('#ctl00_ContentPlaceHolder5_ddlVillage', village)
        time.sleep(1.0)
        # Village selected; moving to matching
        self._emit_status('Matching for Survey Numbers')

        # --- New approach: iterate SubZones rows across pages and validate textarea ---
        print('[Method2] Scanning SubZones rows across pages (no Survey search)...')
        try:
            # Ensure we are on SubZones (default)
            self._log_checked_radio(prefix="Before scan: ")
            if not (self._select_radio_option('SubZones') or self._ensure_radio_selected('SubZones')):
                print('[Method2] Could not explicitly select SubZones; proceeding assuming it is default')
            # Wait for the SubZones grid to be present
            results_table_sel = '#ctl00_ContentPlaceHolder5_dg_Valuation2_0'
            ctx_tbl = self._find_context_with_selector(results_table_sel)
            if ctx_tbl is None:
                return {"error": "SubZones grid not found after village selection"}
            try:
                ctx_tbl.wait_for_selector(results_table_sel, timeout=8000)
            except Exception:
                pass

            # Helper: parse current page rows returning list of tuples (click_row_index, col2_text, col3_rate)
            def _parse_rows():
                try:
                    html_local = ctx_tbl.locator(results_table_sel).inner_html()
                    soup_local = BeautifulSoup(f"<table>{html_local}</table>", 'html.parser')
                    trs = soup_local.find_all('tr')
                    out = []
                    for j, tr in enumerate(trs):
                        if j == 0:
                            continue  # header row
                        tds = tr.find_all('td')
                        if len(tds) < 3:
                            continue
                        # include only rows where first td has a SurveyNo link
                        first_td = tds[0]
                        has_link = first_td.find('a') is not None and ('SurveyNo' in first_td.get_text())
                        if not has_link:
                            continue
                        click_row = j + 1  # nth-child index in live table
                        col2_text = tds[1].get_text(strip=True)
                        col3_rate = tds[2].get_text(strip=True)
                        out.append((click_row, col2_text, col3_rate))
                    return out
                except Exception:
                    return []

            # Helper: signature of the first data row (col2 text) to detect page changes
            def _first_row_sig() -> str:
                rows = _parse_rows()
                if not rows:
                    return ''
                # rows are (click_row, col2, col3)
                return rows[0][1]

            # Searching utility
            def _textbox_has_all_surveys(prev_value: str) -> Tuple[bool, str]:
                try:
                    self.page.wait_for_selector('textarea', timeout=30000)
                except Exception:
                    pass
                # Poll for up to ~30s until textarea value is non-empty and changed
                val = ''
                for _ in range(60):
                    try:
                        # Scroll into view to ensure rendering
                        try:
                            self.page.locator('textarea').first.scroll_into_view_if_needed()
                        except Exception:
                            pass
                        val = self.page.locator('textarea').first.input_value()
                    except Exception:
                        val = ''
                    if val and val.strip() and val != prev_value:
                        break
                    time.sleep(0.5)
                vnorm = (val or '').replace(' ', '')
                needed = [sv.replace(' ', '') for sv in surveys]
                ok = all(sv in vnorm for sv in needed)
                return ok, val or ''

            # Iterate pages until found or exhausted
            current_page = 1
            prev_html = ctx_tbl.locator(results_table_sel).inner_html()
            while True:
                # Small settle before processing a page (esp. first page)
                time.sleep(1.2)
                rows = _parse_rows()
                print(f"[Method2] SubZones: page {current_page} has {len(rows)} data rows (with SurveyNo link)")
                # Try each row on this page
                found = False
                for click_row, col2_text, col3_rate in rows:
                    try:
                        row_css = f"{results_table_sel} tr:nth-child({click_row})"
                        row_loc = ctx_tbl.locator(row_css)
                        link = row_loc.locator("td:nth-child(1) a:has-text('SurveyNo')")
                        if link.count() == 0:
                            link = row_loc.locator("td:nth-child(1) a")
                        if link.count() == 0:
                            continue
                        # Ensure link is visible and scrolled into view
                        try:
                            link.first.scroll_into_view_if_needed()
                        except Exception:
                            pass
                        print(f"[Method2] Clicking SurveyNo at page {current_page}, row {click_row-1}")
                        # Capture previous textarea value, if any
                        try:
                            prev_text_val = self.page.locator('textarea').first.input_value()
                        except Exception:
                            prev_text_val = ''
                        link.first.click()
                        try:
                            self.page.wait_for_load_state('networkidle', timeout=8000)
                        except Exception:
                            pass
                        ok, text_val = _textbox_has_all_surveys(prev_text_val)
                        if ok:
                            print(f"[Method2] All surveys found in textbox for page {current_page}, row {click_row-1}")
                            self._emit_status('Done')
                            try:
                                time.sleep(0.8)
                            except Exception:
                                pass
                            # Print final answer without mentioning column index or surveys list
                            try:
                                print(f"Rate: {col3_rate}")
                            except Exception:
                                pass
                            return {
                                "status": "success",
                                "matched_subzone": col2_text,
                                "rate_value": col3_rate,
                                "textbox_value": text_val,
                            }
                        else:
                            print(f"[Method2] Surveys NOT all present for this row; continuing")
                            # Small pause before the next row to avoid overlapping loads
                            time.sleep(1.0)
                    except Exception:
                        continue

                # Move to next page by clicking next page number inside the grid
                next_page = current_page + 1
                try:
                    pager_link = ctx_tbl.locator(f"{results_table_sel} .cssPager a:has-text('{next_page}')")
                    if pager_link.count() == 0:
                        pager_link = ctx_tbl.get_by_text(str(next_page), exact=True)
                except Exception:
                    pager_link = None
                if not pager_link or pager_link.count() == 0:
                    break
                try:
                    print(f"[Method2] Going to SubZones page {next_page}")
                    # capture current first-row signature before navigation
                    pre_sig = _first_row_sig()
                    pager_link.first.click()
                except Exception:
                    break
                # Wait for grid to change
                for _ in range(60):  # up to ~12s
                    time.sleep(0.2)
                    try:
                        curr_html = ctx_tbl.locator(results_table_sel).inner_html()
                        if curr_html != prev_html:
                            prev_html = curr_html
                            break
                    except Exception:
                        pass
                # Additionally wait for first row signature to change (safer than html diff alone)
                for _ in range(40):  # up to ~8s
                    sig = _first_row_sig()
                    if sig and sig != pre_sig:
                        break
                    time.sleep(0.2)
                time.sleep(1.2)
                current_page = next_page

            return {"error": "Exhausted all SubZones pages and rows without finding all surveys"}
        except Exception as e:
            return {"error": f"Failed during SubZones scan: {e}"}

        # 4) Select radio option 'Survey No.'
        time.sleep(0.5)
        try:
            self.page.wait_for_selector("input[type='radio']", timeout=5000)
        except Exception:
            pass
        # Try robust selection and forced check
        if not (self._select_radio_option('Survey No.') or self._select_radio_option('Survey No') or self._select_radio_option('Survey') or self._ensure_radio_selected('Survey No') or self._ensure_radio_selected('Survey')):
            raise RuntimeError("Could not select 'Survey No.' radio option")
        time.sleep(0.5)
        self._log_checked_radio(prefix="After Survey selection: ")
        # Wait explicitly for the known survey input to appear if present on this page
        try:
            self.page.wait_for_selector('#ctl00_ContentPlaceHolder5_txtCommonSurvey', timeout=5000)
        except Exception:
            pass

        # 5) Enter Survey No: Use only the numeric part of the first survey number
        if not surveys:
            return {"error": "No survey numbers extracted from document"}
        first_numeric = re.match(r"^(\d+)", surveys[0])
        if not first_numeric:
            return {"error": f"First survey number not numeric at start: {surveys[0]}"}
        self._set_input_by_label_text('Enter Survey No', first_numeric.group(1))

        # Click 'Search'
        if not self._click_by_text_any('Search', exact=True):
            if not self._click_by_text_any('Search', exact=False):
                raise RuntimeError("Could not click 'Search'")
        # Allow network/dom to settle
        try:
            self.page.wait_for_load_state('networkidle', timeout=5000)
        except Exception:
            pass
        time.sleep(0.8)

        # 7) Record every row in first column of the resulting table
        print('[Method2] Waiting for first results table...')
        # Prefer Survey No. results grid by exact id discovered in dumps
        survey_results_sel = '#ctl00_ContentPlaceHolder5_grdUrbanRate'
        table_ctx = self._find_context_with_selector(survey_results_sel)
        if table_ctx is not None:
            try:
                table_ctx.wait_for_selector(survey_results_sel, timeout=5000)
            except Exception:
                pass
            try:
                table_html = table_ctx.locator(survey_results_sel).inner_html()
                print('[Method2] Using exact Survey results table by id ctl00_ContentPlaceHolder5_grdUrbanRate')
            except Exception:
                table_ctx = None
                table_html = ''
        else:
            table_html = ''
        # Fallback to heuristic scan if exact id not found
        if not table_html:
            table_ctx, table_html = self._wait_and_get_results_table_html()
        if table_ctx is None or not table_html:
            return {"error": "Could not find first results table"}
        try:
            # Log a short preview of the selected table's text to verify correctness
            tbl_preview = BeautifulSoup(f"<table>{table_html}</table>", 'html.parser').get_text(" ", strip=True)
            print(f"[Method2] Selected table preview (first 200 chars): {tbl_preview[:200]}")
        except Exception:
            pass
        soup = BeautifulSoup(f"<table>{table_html}</table>", 'html.parser')
        recorded_first_col = []
        for tr in soup.find_all('tr')[1:]:
            tds = tr.find_all('td')
            if not tds:
                continue
            recorded_first_col.append(tds[0].get_text(strip=True))
        print(f"[Method2] Recorded {len(recorded_first_col)} entries from first column")
        try:
            print("[Method2] First-column entries:")
            for i, val in enumerate(recorded_first_col, 1):
                print(f"  {i}. {val}")
        except Exception:
            pass

        # 8) Select radio option 'SubZones'
        time.sleep(0.5)
        try:
            self.page.wait_for_selector("input[type='radio']", timeout=5000)
        except Exception:
            pass
        if not (self._select_radio_option('SubZones') or self._ensure_radio_selected('SubZones')):
            return {"error": "Could not switch to SubZones"}
        time.sleep(1.0)
        self._log_checked_radio(prefix="After SubZones selection: ")

        # 9-10) Iterate SubZones grid pages and match where col2 contains any of recorded_first_col (normalized)
        results_table_sel = '#ctl00_ContentPlaceHolder5_dg_Valuation2_0'
        matched_row_info = None

        def _norm_mr(s: str) -> str:
            s = (s or '').strip()
            # remove spaces and common punctuation, unify case
            s = re.sub(r"[\s\-\(\)\.:/]+", "", s)
            return s.lower()

        recorded_norm = [_norm_mr(x) for x in recorded_first_col if x]
        print(f"[Method2] Recorded normalized keys: {recorded_norm}")

        def parse_current_page():
            nonlocal matched_row_info
            ctx_tbl = self._find_context_with_selector(results_table_sel)
            if ctx_tbl is None:
                return False
            html = ctx_tbl.locator(results_table_sel).inner_html()
            soup2 = BeautifulSoup(f"<table>{html}</table>", 'html.parser')
            rows = soup2.find_all('tr')[1:]
            for idx, tr in enumerate(rows):
                tds = tr.find_all('td')
                if len(tds) < 3:
                    continue
                col1 = tds[0].get_text(strip=True)
                col2 = tds[1].get_text(strip=True)
                col3 = tds[2].get_text(strip=True)
                # 10) Prefer exact normalized equality on column 2, fallback to contains
                c2n = _norm_mr(col2)
                matched_key = None
                for k in recorded_norm:
                    if not k:
                        continue
                    if c2n == k:
                        matched_key = k
                        print(f"[Method2] Match equality: row {idx+1} col2='{col2}' norm='{c2n}' == key='{k}'")
                        break
                if matched_key is None:
                    for k in recorded_norm:
                        if not k:
                            continue
                        if (k in c2n) or (c2n in k):
                            matched_key = k
                            print(f"[Method2] Match contains: row {idx+1} col2='{col2}' norm='{c2n}' ~ key='{k}'")
                            break
                if matched_key is not None:
                    matched_row_info = {"row_index": idx + 1, "col1": col1, "col2": col2, "col3": col3, "matched_key": matched_key, "col2_norm": c2n}
                    print(f"[Method2] Selected row {idx+1} -> col2='{col2}', keyUsed='{matched_key}'")
                    return True
            return False

        # Try current page, then paginate slowly ensuring table reloads
        if not parse_current_page():
            try:
                ctx_tbl = self._find_context_with_selector(results_table_sel)
                if ctx_tbl is None:
                    return {"error": "SubZones grid not found"}
                prev_html = ctx_tbl.locator(results_table_sel).inner_html()
                for page_num in range(2, 51):
                    print(f"[Method2] Going to SubZones page {page_num}")
                    moved = False
                    try:
                        # Click page link within the grid pager
                        pager_link = ctx_tbl.locator(f"{results_table_sel} .cssPager a:has-text('{page_num}')")
                        if pager_link.count() == 0:
                            # fallback to any exact text within grid
                            pager_link = ctx_tbl.get_by_text(str(page_num), exact=True)
                        if pager_link.count() > 0:
                            pager_link.first.click()
                            moved = True
                    except Exception:
                        moved = False
                    if not moved:
                        break
                    # Wait for grid to reload by innerHTML changing
                    for _ in range(20):  # up to ~4s
                        try:
                            time.sleep(0.2)
                            curr_html = ctx_tbl.locator(results_table_sel).inner_html()
                            if curr_html != prev_html:
                                prev_html = curr_html
                                break
                        except Exception:
                            time.sleep(0.2)
                    # Give a little extra time for rows to settle
                    time.sleep(0.6)
                    if parse_current_page():
                        break
            except Exception:
                pass

        if not matched_row_info:
            return {"error": "No matching subzone row found"}

        # 11) Click 'SurveyNo' link within the matched row only
        try:
            ctx_tbl_click = self._find_context_with_selector(results_table_sel) or self.page
            row_css = f"{results_table_sel} tr:nth-child({matched_row_info['row_index']+1})"
            row_loc = ctx_tbl_click.locator(row_css)
            # Prefer anchor with exact text SurveyNo inside the first column
            link = row_loc.locator("td:nth-child(1) a:has-text('SurveyNo')")
            if link.count() == 0:
                # Fallback: any anchor in first column
                link = row_loc.locator("td:nth-child(1) a")
            if link.count() == 0:
                return {"error": "SurveyNo link not found in matched row"}
            link.first.click()
        except Exception:
            return {"error": "Could not open SurveyNo details"}

        # 12) After click, wait for the textarea with all survey numbers
        try:
            self.page.wait_for_selector('textarea', timeout=15000)
        except Exception:
            pass
        # Wait until textarea value is populated
        try:
            self.page.wait_for_function("() => { const ta = document.querySelector('textarea'); return ta && ta.value && ta.value.trim().length > 0; }", timeout=15000)
        except Exception:
            pass
        time.sleep(1.0)
        textbox_value = ""
        try:
            textbox_value = self.page.locator('textarea').first.input_value()
        except Exception:
            try:
                textbox_value = self.page.locator("input[type='text']").nth(1).input_value()
            except Exception:
                textbox_value = ""

        if not textbox_value:
            return {"error": "Could not read survey numbers textbox"}

        # 13) Verify that all considered surveys exist in the textbox
        text_norm = textbox_value.replace(' ', '')
        all_present = all(sv.replace(' ', '') in text_norm for sv in surveys)
        if not all_present:
            return {"error": "Not all input survey numbers are present in the area details"}

        # 14) If all exist, use column 3 value for that matched row
        return {
            "status": "success",
            "matched_subzone": matched_row_info["col2"],
            "rate_value": matched_row_info["col3"],
            "surveys_checked": surveys,
        }


# ----------------------
# Public API
# ----------------------

def process_igr_from_doc(file_bytes: bytes, filename: str, year_label: str, district_override: Optional[str] = None, taluka_override: Optional[str] = None, village_override: Optional[str] = None, progress_cb: Optional[Callable[[str], None]] = None) -> Dict:
    """
    Process a Word/PDF file and fetch rate based on IGR SubZones matching.
    - filename: used to branch parsing logic. Currently supports .docx only.
    - year_label: dropdown label like '2015-2016' ... '2025-2026'.
    Returns dict with success or error.
    """
    name_lower = (filename or '').lower()
    if name_lower.endswith('.docx'):
        district, taluka, village, surveys = extract_admin_and_surveys_from_docx(file_bytes)
    else:
        return {"error": "Only .docx is supported currently for Method 2"}

    # Apply user overrides when provided
    district = district_override or district
    taluka = taluka_override or taluka
    village = village_override or village

    if not district or not taluka or not village or not surveys:
        return {"error": "Could not extract district/taluka/village/surveys from the document"}

    with IGRSubzoneScraper(headless=False, progress_cb=progress_cb) as scraper:
        return scraper.run(district=district, year_label=year_label, taluka=taluka, village=village, surveys=surveys, translate_admin=False)
