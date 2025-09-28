import os
import pathlib
import json
import re
import csv
import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from dotenv import load_dotenv
import google.generativeai as genai
import math

# ----------------------
# Configuration & Paths
# ----------------------
PDF_SOURCE_FOLDER = "index2/files"
CSV_EXPORT_DIR = "index2/OCR_CSVexports"
CSV_OUTPUT_FILE = os.path.join(CSV_EXPORT_DIR, "output.csv")
WORD_TEMPLATE_FILE = "index2/format.docx"
WORD_EXPORT_DIR = "index2/OCR_WORDexports"
WORD_OUTPUT_FILE = os.path.join(WORD_EXPORT_DIR, "output.docx")

# Model name is hardcoded per requirement
GEMINI_MODEL_NAME = "models/gemini-2.5-pro"


def load_api_key():
    """Load GOOGLE_API_KEY from .env and configure the Gemini client."""
    load_dotenv()  # loads from .env in project root
    api_key = os.getenv("GOOGLE_API_KEY")
    if not api_key:
        raise RuntimeError("GOOGLE_API_KEY not found. Create a .env with GOOGLE_API_KEY=<your_key>.")
    genai.configure(api_key=api_key)


def get_multipage_extraction_prompt():
    """Returns the detailed prompt for the Gemini model to handle multi-page PDFs."""
    return """
    Analyze the entire provided multi-page Marathi PDF document. Find ALL pages that are
    formatted as 'Index-II' (सूची क्र.2).

    **CRITICAL INSTRUCTIONS:**
    1.  **Identify and Process ONLY 'Index-II' pages.** These pages contain numbered
        fields like '(1)विलेखाचा प्रकार', '(5)क्षेत्रफळ', etc.
    2.  **Explicitly IGNORE irrelevant pages.** Completely skip any pages titled
        'Payment Details' or any pages that primarily consist of a table of financial
        transactions. Do not extract any data from them.
    3.  **Return a JSON array.** The final output must be a single JSON array (a list
        of objects), where each JSON object represents the data extracted from ONE
        'Index-II' page.
    4.  If no 'Index-II' pages are found, return an empty array `[]`.

    **For each 'Index-II' page found, extract the following fields into a JSON object:**
    -   `dast_kramank_year`: The year part of 'दस्त क्रमांक'.
    -   `sub_registrar_number`: The number at the end of 'दुय्यम निबंधक'.
    -   `dast_kramank_full`: The full value of 'दस्त क्रमांक'.
    -   `registration_date`: The value for '(10)दस्त नोंदणी केल्याचा दिनांक'.
    -   `document_type`: The value for '(1)विलेखाचा प्रकार'.
    -   `survey_number`: All the numbers inside the double parentheses `((...))` from section (4),
        like 'Survey Number'. Some numbers can have parts to them (for example 1ब, 2ब)
    -   `area_sq_meter`: The value for '(5)क्षेत्रफळ'. VERY IMPORTANT: If the unit is
        'चौ.फुट' (Square Feet) or 'चौ. फूट', convert it to Square Meters by multiplying by
        0.092903. If the unit is 'हेक्टर' (Hectare), convert to Square Meters by multiplying
        by 10,000. If it is already 'चौ.मीटर', use the value directly. Return only the
        final numerical value in Square Meters.
    -   `stamp_duty`: The value for '(12)बाजारभावाप्रमाणे मुद्रांक शुल्क'.
    -   `prakar`: Analyze the text in section '(4) भू-मापन...' and related context. If it contains
        the Marathi word 'सदनिका' or mentions चौ.फूट, set `prakar` to 'सदनिका'. Otherwise set it to
        'बिनशेती जमिन'.
    -   `amount`: The value for '(2) मोबदला'. Return only the numeric value (no currency symbols).

    Return ONLY the JSON array and nothing else.
    """


def clean_and_convert_to_float(value, default=0.0):
    """Safely converts a value to a float."""
    if value is None:
        return default
    try:
        cleaned_value = re.sub(r"[^0-9.]", "", str(value))
        return float(cleaned_value) if cleaned_value else default
    except (ValueError, TypeError):
        return default


def process_multipage_pdf(pdf_path: pathlib.Path, model, prompt: str):
    """Processes a single multi-page PDF, which may contain multiple records.
    Returns list of processed record dicts.
    """
    print(f"-> Uploading and processing '{pdf_path.name}'...")
    all_records = []
    try:
        pdf_file = genai.upload_file(path=pdf_path, display_name=pdf_path.name)
        response = model.generate_content([prompt, pdf_file])
        json_text = response.text.strip().replace("```json", "").replace("```", "")
        list_of_data = json.loads(json_text)

        if not isinstance(list_of_data, list):
            print(f"   ...Warning: Expected a list of records for '{pdf_path.name}', but did not receive one.")
            return []

        print(f"   ...Found {len(list_of_data)} relevant record(s) in '{pdf_path.name}'.")

        for i, data in enumerate(list_of_data):
            area_sqm = clean_and_convert_to_float(data.get("area_sq_meter"))
            stamp_duty = clean_and_convert_to_float(data.get("stamp_duty"))
            amount = clean_and_convert_to_float(data.get("amount"))

            hectares = area_sqm / 10000 if area_sqm > 0 else 0
            rate_per_sqm = stamp_duty / area_sqm if area_sqm > 0 else 0
            # Approximation per user's choice: 1 guntha ~ 100 sqm
            rate_per_guntha = rate_per_sqm * 100 if rate_per_sqm > 0 else 0
            rate_per_ha = rate_per_guntha * 100 if rate_per_guntha > 0 else 0  # = rate_per_sqm * 10000

            processed_data = {
                "dast_kramank_year": data.get("dast_kramank_year", "N/A"),
                "sub_registrar_number": data.get("sub_registrar_number", "N/A"),
                "dast_kramank_full": data.get("dast_kramank_full", "N/A"),
                "registration_date": data.get("registration_date", "N/A"),
                "document_type": data.get("document_type", "N/A"),
                "survey_number": data.get("survey_number", "N/A"),
                "area_sq_meter": f"{area_sqm:.4f}",
                "area_hectares": f"{hectares:.8f}",
                "stamp_duty": f"{stamp_duty:.2f}",
                "rate_per_sqm": f"{rate_per_sqm:.2f}",
                "rate_per_guntha": f"{rate_per_guntha:.2f}",
                "rate_per_ha": f"{rate_per_ha:.2f}",
                "prakar": data.get("prakar", "N/A"),
                "amount": f"{amount:.2f}",
                "source_file": pdf_path.name,
                "page_record_num": i + 1,
            }
            all_records.append(processed_data)

        return all_records

    except Exception as e:
        print(f"   ...An error occurred while processing {pdf_path.name}: {e}")
        return []


def export_csv(records_by_file, csv_output: str):
    """Export all records to CSV (overwrite existing)."""
    os.makedirs(os.path.dirname(csv_output), exist_ok=True)

    headers = [
        "Serial Number", "(2) Year", "(3) Sub Registrar Number", "(4) Dast Kramank",
        "(5) Registration Date", "(6) Document Type", "(7) Survey Number",
        "(8) Area (sq meters)", "(9) Area (Hectares)", "(10) Stamp Duty",
        "(11) Rate per SqM", "(12) Rate per Guntha", "(13) Rate per Ha",
        "प्रकार", "Amount", "Source PDF", "Record # in PDF"
    ]

    with open(csv_output, "w", newline="", encoding="utf-8") as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(headers)

        serial_number = 1
        for rec in records_by_file:
            row = [
                serial_number, rec["dast_kramank_year"], rec["sub_registrar_number"], rec["dast_kramank_full"],
                rec["registration_date"], rec["document_type"], rec["survey_number"],
                rec["area_sq_meter"], rec["area_hectares"], rec["stamp_duty"],
                rec["rate_per_sqm"], rec["rate_per_guntha"], rec["rate_per_ha"],
                rec.get("prakar", "N/A"), rec.get("amount", ""), rec["source_file"], rec["page_record_num"]
            ]
            writer.writerow(row)
            serial_number += 1

    print(f"CSV exported to: {csv_output}")


def add_table_borders(table):
    """Add black borders to a python-docx table."""
    tbl = table._element
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border_el = OxmlElement(f"w:{border_name}")
        border_el.set(qn("w:val"), "single")
        border_el.set(qn("w:sz"), "8")
        border_el.set(qn("w:space"), "0")
        border_el.set(qn("w:color"), "000000")
        tblBorders.append(border_el)
    tbl.tblPr.append(tblBorders)


def export_word_from_csv(csv_path: str, template_path: str, output_path: str):
    """Populate the first table of the template with CSV data and save to output_path.
    It mirrors the existing csv-table-process.py behavior (dropping last two columns).
    """
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    df = pd.read_csv(csv_path)
    # Drop last two columns (e.g., provenance columns)
    if df.shape[1] >= 2:
        df = df.iloc[:, :-2]

    doc = Document(template_path)
    table = doc.tables[0]

    # Append each DataFrame row to the Word table
    for _, row in df.iterrows():
        cells = table.add_row().cells
        # Write up to the number of columns available in the template table
        max_cols = min(len(cells), len(row))
        if len(row) > len(cells):
            # Warn once per row if template has fewer columns than data
            # e.g., template has 14 columns but data has 15 (includes 'Amount')
            print(
                f"Warning: template has {len(cells)} columns but data row has {len(row)}; extra columns will be truncated."
            )
        for i, value in enumerate(row[:max_cols]):
            cells[i].text = "" if pd.isna(value) else str(value)

    # Add borders
    add_table_borders(table)

    # Save (overwrite)
    doc.save(output_path)
    print(f"Word document exported to: {output_path}")


def main():
    # Prep directories
    os.makedirs(CSV_EXPORT_DIR, exist_ok=True)
    os.makedirs(WORD_EXPORT_DIR, exist_ok=True)
    pathlib.Path(PDF_SOURCE_FOLDER).mkdir(exist_ok=True)

    # Configure API
    load_api_key()

    # Prepare model
    model = genai.GenerativeModel(model_name=GEMINI_MODEL_NAME)

    # Find PDFs
    source_path = pathlib.Path(PDF_SOURCE_FOLDER)
    pdf_files = list(source_path.glob("*.pdf"))
    if not pdf_files:
        print(f"No PDF files found in '{PDF_SOURCE_FOLDER}'. Place files and run again.")
        return

    print(f"Found {len(pdf_files)} PDF file(s) to process.")

    # Process PDFs
    prompt = get_multipage_extraction_prompt()
    all_records = []
    for pdf_path in pdf_files:
        records_from_file = process_multipage_pdf(pdf_path, model, prompt)
        all_records.extend(records_from_file)

    # Export CSV (overwrite)
    export_csv(all_records, CSV_OUTPUT_FILE)

    # Export Word from CSV (overwrite)
    export_word_from_csv(CSV_OUTPUT_FILE, WORD_TEMPLATE_FILE, WORD_OUTPUT_FILE)
    
    # Post-processing on the generated Word file
    future_filter_and_aggregate(WORD_OUTPUT_FILE)

def future_filter_and_aggregate(word_file_path: str):
    """Filter rows by 'प्रकार' == 'बिनशेती जमिन', create a new table under the previous
    table with same formatting, sort by 'प्रती चौ.मी.' descending, keep top 50% (rounding up),
    compute average of 'प्रती चौ.मी.' and append it. Assumes first table has a merged title row,
    header row at index 1, number row at index 2, and data from index 3 onwards."""

    doc = Document(word_file_path)

    if not doc.tables:
        print("No tables found in the document for future processing.")
        return

    base_table = doc.tables[0]

    # Build a list of rows (as lists of text) and detect column indices
    rows_text = []
    for r_idx, row in enumerate(base_table.rows):
        cell_texts = [cell.text.strip() for cell in row.cells]
        rows_text.append(cell_texts)

    if len(rows_text) < 2:
        print("Base table does not contain enough rows for header/data.")
        return

    # Use row index 1 as the visual header row, but determine column indices by
    # the known CSV export order (headers defined in export_csv) to avoid label mismatches.
    visual_header = rows_text[1]
    kept_headers = [
        "Serial Number", "(2) Year", "(3) Sub Registrar Number", "(4) Dast Kramank",
        "(5) Registration Date", "(6) Document Type", "(7) Survey Number",
        "(8) Area (sq meters)", "(9) Area (Hectares)", "(10) Stamp Duty",
        "(11) Rate per SqM", "(12) Rate per Guntha", "(13) Rate per Ha",
        "प्रकार", "Amount"
    ]
    # Determine column indices by position in kept_headers
    try:
        col_idx_prakar = kept_headers.index("प्रकार")
        col_idx_prati_chou_mi = kept_headers.index("(11) Rate per SqM")
    except ValueError:
        print("Internal header mapping failed; please verify export_csv headers.")
        return

    data_rows = rows_text[3:]

    # Filter by 'प्रकार' == 'बिनशेती जमिन'
    filtered = [
        r for r in data_rows
        if len(r) > max(col_idx_prakar, col_idx_prati_chou_mi) and r[col_idx_prakar] == 'बिनशेती जमिन'
    ]

    # Convert 'प्रती चौ.मी.' to float for sorting
    def to_float_safe(s):
        s_clean = re.sub(r"[^0-9.]", "", s or "")
        try:
            return float(s_clean) if s_clean else 0.0
        except ValueError:
            return 0.0

    filtered.sort(key=lambda r: to_float_safe(r[col_idx_prati_chou_mi]), reverse=True)

    # Keep top 50% rows (round up, at least 1 if any exist)
    n = len(filtered)
    keep_n = max(1, math.ceil(n * 0.5)) if n > 0 else 0
    top_half = filtered[:keep_n]

    # Insert a heading and a new table for filtered rows under the original table
    doc.add_paragraph("बिनशेती जमिन - फिल्टर केलेले")
    new_table = doc.add_table(rows=1, cols=len(visual_header))
    new_table.style = base_table.style
    # Copy header
    for c_idx, text in enumerate(visual_header):
        new_table.rows[0].cells[c_idx].text = text
    # Add rows
    for r in filtered:
        cells = new_table.add_row().cells
        for c_idx, text in enumerate(r[:len(visual_header)]):
            cells[c_idx].text = text
    add_table_borders(new_table)

    # Insert a second heading and third table with top 50%
    doc.add_paragraph("बिनशेती जमिन - टॉप 50% (प्रती चौ.मी. नुसार)")
    top_table = doc.add_table(rows=1, cols=len(visual_header))
    top_table.style = base_table.style
    for c_idx, text in enumerate(visual_header):
        top_table.rows[0].cells[c_idx].text = text
    for r in top_half:
        cells = top_table.add_row().cells
        for c_idx, text in enumerate(r[:len(visual_header)]):
            cells[c_idx].text = text
    add_table_borders(top_table)

    # Compute average of 'प्रती चौ.मी.' from top_half
    if top_half:
        avg_value = sum(to_float_safe(r[col_idx_prati_chou_mi]) for r in top_half) / len(top_half)
    else:
        avg_value = 0.0

    doc.add_paragraph(f"Average प्रती चौ.मी. = {avg_value:.2f}")

    # Overwrite the same output file
    doc.save(word_file_path)


if __name__ == "__main__":
    main()
